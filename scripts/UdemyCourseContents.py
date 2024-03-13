import sys 
from seleniumbase import Driver
from selenium.webdriver.common.by import By
import time
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import logging
from datetime import datetime
import re

#----------------------------------------------------------------------------------------------
# Reading Udemy account credentials and output folder path from the config file.

file_path = 'scripts/config.ini'
with open(file_path, 'r', encoding='utf-8') as file:
    config_content = file.readlines()
    
first_line = config_content[0]
second_line = config_content[1]

username = first_line.split('=')[1].strip()
password = second_line.split('=')[1].strip()

#----------------------------------------------------------------------------------------------
# Reading the command-line argument

if len(sys.argv) < 2:
    sys.exit("Usage: python script.py <SrcFolder> <DestFolder>")

# Get the value from the command-line arguments
course_url = sys.argv[1]

# Check if the second argument is provided, otherwise use a default value
if len(sys.argv) > 2:
    outputs_path = sys.argv[2]
    logs_path = "logs"
    try :
        os.mkdir(outputs_path)
        os.mkdir(logs_path)
        
    except Exception as e:
        pass
else:
    third_line = config_content[2]
    output_folder = third_line.split('=')[1].strip()
    folder_path = output_folder.replace("\\", "/")
    logs_path = os.path.join(folder_path, "logs")
    try :
        outputs_path = os.path.join(folder_path, "output")
        os.mkdir(outputs_path)
        
        os.mkdir(logs_path)
    except Exception as e:
        pass
  
#----------------------------------------------------------------------------------------------
# Validating Udemy URL

def is_valid_udemy_url(course_url):
    udemy_prefix = 'https://www.udemy.com'
    
    if course_url.startswith(udemy_prefix):
        return True
    else:
        return False

#----------------------------------------------------------------------------------------------
# Creating a 'logs' folder and a log file

current_date_time = datetime.now()
formatted_date_time = current_date_time.strftime("%Y-%m-%d %H_%M_%S")

pattern = r'https://www.udemy.com/course/([^/]+)/?$'
match = re.match(pattern, course_url)
l_name = match.group(1)

log_file_name = f'{l_name}-{formatted_date_time}.log'

logging.basicConfig( 
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    filename=os.path.join(logs_path, log_file_name),  
    filemode="w" 
)

#----------------------------------------------------------------------------------------------
# Validating Udemy URL and logging the result

if is_valid_udemy_url(course_url):
    logging.info(f'The URL "{course_url}" is valid.')
else:
    logging.error(f'The URL "{course_url}" is not valid.')
    sys.exit()

#----------------------------------------------------------------------------------------------
# Starting the web driver

driver = Driver(uc=True)


try:
    driver.get(course_url)
    print("starting webdriver service")
    logging.info("starting webdriver service")
    
except Exception as e:
    logging.error (e)
    sys.exit()
    
time.sleep(10)

logging.info("Enrolling the course")
print("Enrolling the course")

#----------------------------------------------------------------------------------------------
# Clicking the Enroll button

try:
    btn1 = driver.find_element(
        By.XPATH, '//*[@id="udemy"]/div[1]/div[2]/div/div/div[2]/div[3]/div[2]/div[2]/div/div/div/div/div[4]/div/button')
    btn1.click()
except Exception as e:
    print (e)
    
time.sleep(5)

logging.info("click log in button")
print("click log in button")

#----------------------------------------------------------------------------------------------
# Clicking the login button

btn2 = driver.find_element(
    By.XPATH, '//*[@id="udemy"]/div[1]/div[1]/div[3]/div[6]/a')
btn2.click()

time.sleep(5)

logging.info("please wait logging......")
print("please wait logging......")

#----------------------------------------------------------------------------------------------
# Entering Udemy credentials and logging in

driver.find_element(By.NAME, "email").send_keys(username)
driver.find_element(By.NAME, "password").send_keys(password)

button = driver.find_element(
    By.XPATH, '//*[@id="udemy"]/div[1]/div[2]/div/div/form/button')
button.click()

logging.info("please wait loading......")
print("please wait loading......")
time.sleep(20)
url = []

#----------------------------------------------------------------------------------------------
# Extracting course information

title = driver.get_title()

url1 = driver.current_url
url.append(url1)

Course_name = f"course name :  {title[8:-5]}"
Course_url = f"course url :  {url1}"
course_content = []

course_content.append(Course_name)
course_content.append(Course_url)

logging.info("Expanding All Section")
print("Expanding All Section")

#----------------------------------------------------------------------------------------------

try:
    driver.click(f'//*[@id="ct-sidebar-scroll-container"]/div/div/div[1]/div[2]/div/ul/li[1]/div')
    time.sleep(2)
except Exception as e:
    driver.click(f'//*[@id="ct-sidebar-scroll-container"]/div/div/div[1]/div[1]')
    time.sleep(2)
    driver.click(f'//*[@id="ct-sidebar-scroll-container"]/div/div/div[1]/div[2]/div/ul/li[1]/div')
    time.sleep(2)
    print("please rerun program")
    driver.quit()
    quit()
    
#----------------------------------------------------------------------------------------------    
# Expanding course sections and Extracting URLs of each section

for x in range(2,15):
    try:
        driver.click(f'//*[@id="ct-sidebar-scroll-container"]/div/div/div[{x}]/div[1]')
        time.sleep(2)
        driver.click(f'//*[@id="ct-sidebar-scroll-container"]/div/div/div[{x}]/div[2]/div/ul/li[1]/div')
        time.sleep(2)
        u = driver.current_url
        url.append(u)
    except Exception as e:
        driver.click(f'//*[@id="ct-sidebar-scroll-container"]/div/div/div[1]/div[2]/div/ul/li[1]/div')
        time.sleep(2)
        u = driver.current_url
        url.append(u)
        break    
#----------------------------------------------------------------------------------------------    
# Extracting HTML content

html_content = driver.page_source

time.sleep(5)
driver.quit()

#----------------------------------------------------------------------------------------------
# Parsing HTML content using BeautifulSoup

soup = BeautifulSoup(html_content, 'html.parser')
course_titel = soup.find('h2', class_='ud-heading-md')
course_content_elements = soup.find_all('span', class_='truncate-with-tooltip--ellipsis--YJw4N')

name = []
for index in course_content_elements:
    name.append(str(index.text))

for i in course_titel:
    course_content.append(str(i.text))
    
for z in course_content_elements:
    content_only = z.find_all('span')
    for item in content_only:
        content = item.text
        course_content.append(str(content))
    
course_content_strings  = '\n'.join(course_content)

folder_path = outputs_path

try:
    directory_name = title[8:-5]
    directory_name_no_spaces = directory_name.replace(" ", "")
    directory_name_no_pipe = directory_name_no_spaces.replace("|", "")
    full_path = os.path.join(folder_path, directory_name_no_pipe)
    os.mkdir(full_path)
    print(f"Directory '{directory_name_no_pipe}' successfully create.")
    logging.info(f"Directory '{directory_name_no_pipe}' successfully create.")
except FileExistsError:
    print(f"Directory '{directory_name_no_pipe}' already exists.")
    logging.info(f"Directory '{directory_name_no_pipe}' already exists.")
    
f_path = f"{folder_path}/{directory_name_no_pipe}"
file_name = '00-readme.txt'
file_path = os.path.join(f_path, file_name)

with open(file_path, "w", encoding="utf-8") as new_file:
    new_file.write(str(course_content_strings))

print(f"File '{file_name}' created inside the folder '{folder_path}'.")
logging.info(f"File '{file_name}' created inside the folder '{folder_path}'.")

#----------------------------------------------------------------------------------------------
# Extracting section-wise content

result_arrays = []
current_array = []

for line in name:
    if line.startswith('Section'):
        if current_array:
            result_arrays.append(current_array)
        current_array = [line]
    else:
        current_array.append(line)

if current_array:
    result_arrays.append(current_array)
    
#----------------------------------------------------------------------------------------------       
# Processing section content and creating Word documents

file_path = 'scripts/00-template.docx'

for idx, section_list in enumerate(result_arrays, start=1):
    try:
        section_name = section_list[0]
        capitalized_name = section_name.title()
        directory_name_no_spaces = capitalized_name.replace(" ", "")
        directory_name_no_pipe = directory_name_no_spaces.replace("Section", "s")
        directory_name_no_flash = directory_name_no_pipe.replace("/", "")
        directory_name_no_pipe_change = directory_name_no_flash.replace(":", "-")
        full_path = os.path.join(f_path, f'{directory_name_no_pipe_change}.docx')
        
        # Create a new document using the template file
        ex_doc = Document(file_path)

        margin_top = ex_doc.sections[0].top_margin.inches
        margin_bottom = ex_doc.sections[0].bottom_margin.inches
        margin_left = ex_doc.sections[0].left_margin.inches
        margin_right = ex_doc.sections[0].right_margin.inches
        
        # Set margins for the new document
        # Set margins for the new document using the values from the existing document
        document = Document()
        section = document.sections[0]
        section.left_margin = Inches(margin_top)
        section.right_margin = Inches(margin_bottom)
        section.top_margin = Inches(margin_left)
        section.bottom_margin = Inches(margin_right)  
        
        # Add section URLs
        document.add_paragraph(f"Section URL: {url[idx - 1]}")

        # Add section contents
        document.add_paragraph("Contents of Section:")
        for item in section_list[1:]:
            document.add_paragraph(item)
            document.add_paragraph("*" * 100)

        document.save(full_path)
        print(f"Word document {directory_name_no_pipe_change}.docx successfully created.")
        logging.info(f"Word document {directory_name_no_pipe_change}.docx successfully created.")
    except Exception as e:
        print(f"Error creating Word document: {e}")
        logging.info(f"Error creating Word document: {e}")
        
#----------------------------------------------------------------------------------------------       
